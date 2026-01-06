from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import textwrap

def to_markdown_bytes(text: str) -> bytes:
    return text.encode("utf-8")

def to_docx_bytes(title: str, text: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.splitlines():
        doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def to_pdf_bytes(title: str, text: str) -> bytes:
    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=letter)
    width, height = letter

    x = 50
    y = height - 60
    line_height = 14

    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, title)
    y -= 24
    c.setFont("Helvetica", 11)

    for raw_line in text.splitlines():
        raw_line = raw_line.rstrip()
        if not raw_line:
            y -= line_height
            continue

        wrapped = textwrap.wrap(raw_line, width=95)
        for wline in wrapped:
            if y < 60:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - 60
            c.drawString(x, y, wline)
            y -= line_height

    c.save()
    return bio.getvalue()
